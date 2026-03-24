import docx
import sys
import codecs

with codecs.open("tables.txt", "w", "utf-8") as f:
    try:
        doc = docx.Document('Reuniones Peripallozza..docx')
        f.write("=== PARAGRAPHS ===\n")
        for para in doc.paragraphs:
            if para.text.strip():
                f.write(para.text + "\n")
                
        f.write("\n=== TABLES ===\n")
        for i, table in enumerate(doc.tables):
            f.write(f"\n--- Table {i} ---\n")
            for row in table.rows:
                row_data = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
                f.write(" | ".join(row_data) + "\n")
    except Exception as e:
        f.write(f"Docx error: {e}\n")
